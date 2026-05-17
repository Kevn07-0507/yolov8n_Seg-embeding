"""
生成实施方案报告 — 含真正框图 + .docx 输出
"""
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.font_manager as fm
import numpy as np
import os

# 设置中文字体
font_paths = fm.findSystemFonts()
cn_fonts = [f for f in fm.fontManager.ttflist if 'Microsoft YaHei' in f.name or 'SimHei' in f.name]
if cn_fonts:
    plt.rcParams['font.family'] = cn_fonts[0].name
else:
    # 尝试直接指定常见中文字体
    for fname in ['Microsoft YaHei', 'SimHei', 'STXihei', 'KaiTi']:
        try:
            plt.rcParams['font.family'] = fname
            break
        except:
            pass
plt.rcParams['axes.unicode_minus'] = False

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = "报告"
os.makedirs(OUT_DIR, exist_ok=True)

# ═══════════════════════════════════════════════════════════
# 框图绘制函数
# ═══════════════════════════════════════════════════════════

def draw_system_architecture():
    """绘制系统总体架构框图"""
    fig, ax = plt.subplots(1, 1, figsize=(16, 8))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_facecolor('#F5F7FA')

    colors = {
        'data': '#4A90D9',
        'train': '#50B86C',
        'infer': '#E8913A',
        'app': '#9B59B6',
        'deploy': '#E74C3C',
        'light': '#ECF0F1',
    }

    def draw_box(ax, x, y, w, h, text, color, fontcolor='white', fontsize=11, bold=True):
        box = FancyBboxPatch((x, y), w, h,
                             boxstyle="round,pad=0.15",
                             facecolor=color, edgecolor='#2C3E50', linewidth=1.5)
        ax.add_patch(box)
        ax.text(x + w/2, y + h/2, text, ha='center', va='center',
                fontsize=fontsize, fontweight='bold' if bold else 'normal',
                color=fontcolor)

    def draw_sub_box(ax, x, y, w, h, text, color, fontsize=9):
        box = FancyBboxPatch((x, y), w, h,
                             boxstyle="round,pad=0.08",
                             facecolor=color, edgecolor='#BDC3C7', linewidth=1)
        ax.add_patch(box)
        ax.text(x + w/2, y + h/2, text, ha='center', va='center',
                fontsize=fontsize, color='#2C3E50')

    def draw_arrow(ax, x1, y1, x2, y2, style='->'):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle=style, color='#7F8C8D', lw=2))

    # 标题
    ax.text(8, 7.6, '墙面裂缝智能巡检系统 — 总体架构', ha='center', va='center',
            fontsize=18, fontweight='bold', color='#2C3E50')

    # ── 五大模块 ──
    # 数据层
    draw_box(ax, 0.3, 5.2, 2.8, 1.6, '数据层\n(Data Layer)', colors['data'])
    draw_sub_box(ax, 0.5, 5.4, 1.2, 0.6, 'crack-seg\n数据集', colors['light'])
    draw_sub_box(ax, 1.85, 5.4, 1.1, 0.6, '数据增强\n预处理', colors['light'])
    draw_sub_box(ax, 0.5, 6.15, 2.45, 0.45, 'Train 3717 | Valid 112 | Test 200', colors['light'])

    # 训练层
    draw_box(ax, 3.5, 5.2, 2.8, 1.6, '训练层\n(Training Layer)', colors['train'])
    draw_sub_box(ax, 3.7, 5.4, 1.2, 0.6, 'YOLOv8-seg\n模型训练', colors['light'])
    draw_sub_box(ax, 5.05, 5.4, 1.1, 0.6, '100 Epochs\n验证与早停', colors['light'])
    draw_sub_box(ax, 3.7, 6.15, 2.45, 0.45, 'mAP / Loss 曲线监控', colors['light'])

    # 推理层
    draw_box(ax, 6.7, 5.2, 2.8, 1.6, '推理层\n(Inference Layer)', colors['infer'])
    draw_sub_box(ax, 6.9, 5.4, 1.2, 0.6, '单张/批量\n图片检测', colors['light'])
    draw_sub_box(ax, 8.25, 5.4, 1.1, 0.6, '实时视频\n摄像头检测', colors['light'])
    draw_sub_box(ax, 6.9, 6.15, 2.45, 0.45, 'NMS + 置信度过滤', colors['light'])

    # 应用层
    draw_box(ax, 9.9, 5.2, 2.8, 1.6, '应用层\n(Application Layer)', colors['app'])
    draw_sub_box(ax, 10.1, 5.4, 1.2, 0.6, 'CLI 命令行\nmain.py', colors['light'])
    draw_sub_box(ax, 11.45, 5.4, 1.1, 0.6, 'Web 界面\nGradio', colors['light'])
    draw_sub_box(ax, 10.1, 6.15, 2.45, 0.45, '可视化 + 报告生成', colors['light'])

    # 部署层
    draw_box(ax, 13.1, 5.2, 2.6, 1.6, '部署层\n(Deploy Layer)', colors['deploy'])
    draw_sub_box(ax, 13.25, 5.4, 1.1, 0.4, 'ONNX', colors['light'])
    draw_sub_box(ax, 13.25, 5.95, 1.1, 0.4, 'TensorRT', colors['light'])
    draw_sub_box(ax, 14.5, 5.4, 1.05, 0.4, 'OpenVINO', colors['light'])
    draw_sub_box(ax, 14.5, 5.95, 1.05, 0.4, 'TFLite', colors['light'])
    draw_sub_box(ax, 13.25, 6.5, 2.3, 0.45, 'Docker 容器 / 边缘设备', colors['light'])

    # 层间箭头
    for x1, x2 in [(3.1, 3.5), (6.3, 6.7), (9.5, 9.9), (12.7, 13.1)]:
        draw_arrow(ax, x1, 6.0, x2, 6.0)

    # ── 流程箭头（底部） ──
    y_flow = 4.5
    flow_steps = [
        ('输入图像', 1.5, colors['data']),
        ('预处理\nResize 640×640', 4.5, colors['data']),
        ('Backbone\nCSPDarknet', 7.5, colors['train']),
        ('Neck\nPAN-FPN', 10.5, colors['train']),
        ('Head\nSegmentation', 13.5, colors['infer']),
    ]
    for i, (text, x, c) in enumerate(flow_steps):
        draw_box(ax, x-1.0, 3.5, 2.0, 0.85, text, c, fontsize=9)
        if i < len(flow_steps) - 1:
            draw_arrow(ax, x + 1.0, 3.93, x + 1.8, 3.93)

    # 输出标注
    ax.text(8, 3.0, '输出: Bounding Box + Segmentation Mask + Confidence\n→ 裂缝面积占比 → 严重程度分级 (轻微/中等/严重)',
            ha='center', va='center', fontsize=10, color='#34495E',
            bbox=dict(boxstyle='round,pad=0.5', facecolor='#FDEBD0', edgecolor='#E8913A'))

    # 图例
    legend_y = 2.2
    for i, (label, c) in enumerate([
        ('数据层', colors['data']), ('训练层', colors['train']),
        ('推理层', colors['infer']), ('应用层', colors['app']), ('部署层', colors['deploy'])
    ]):
        ax.add_patch(FancyBboxPatch((1 + i*3, legend_y), 2.5, 0.4,
                     boxstyle="round,pad=0.05", facecolor=c, edgecolor='#2C3E50'))
        ax.text(2.25 + i*3, legend_y + 0.2, label, ha='center', va='center', fontsize=9, color='white', fontweight='bold')

    plt.tight_layout()
    path = os.path.join(OUT_DIR, "fig1_system_architecture.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def draw_inference_pipeline():
    """绘制单张图像推理处理流程"""
    fig, ax = plt.subplots(1, 1, figsize=(14, 4.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 4.5)
    ax.axis('off')
    ax.set_facecolor('#FAFBFC')

    colors = ['#3498DB', '#2ECC71', '#E67E22', '#9B59B6', '#E74C3C', '#1ABC9C', '#F39C12']

    steps = [
        ('① 输入\n墙面图像', 'Raw Image\n.jpg / .png'),
        ('② 预处理\nResize + Normalize', '640×640×3\nTensor'),
        ('③ Backbone\nCSPDarknet', '多尺度\n特征图'),
        ('④ Neck\nPAN-FPN', '特征金字塔\n融合'),
        ('⑤ Head\nSeg Detect', 'BBox + Mask\n+ Confidence'),
        ('⑥ NMS\n后处理', '去重 + 置信度\n过滤'),
        ('⑦ 输出\n可视化', '标注图 + \n面积 + 严重度'),
    ]

    box_w = 1.6
    gap = 0.3
    y_top = 2.8
    y_sub = 0.8

    for i, (title, desc) in enumerate(steps):
        x = 0.5 + i * (box_w + gap)
        c = colors[i]
        # 主框
        box = FancyBboxPatch((x, y_top), box_w, 0.9,
                             boxstyle="round,pad=0.12", facecolor=c, edgecolor='#2C3E50', linewidth=1.5)
        ax.add_patch(box)
        ax.text(x + box_w/2, y_top + 0.45, title, ha='center', va='center',
                fontsize=8, fontweight='bold', color='white')
        # 子框
        sub = FancyBboxPatch((x, y_sub), box_w, 0.65,
                             boxstyle="round,pad=0.08", facecolor='#FFFFFF', edgecolor=c, linewidth=1.2)
        ax.add_patch(sub)
        ax.text(x + box_w/2, y_sub + 0.33, desc, ha='center', va='center', fontsize=7, color='#2C3E50')
        # 箭头
        if i < len(steps) - 1:
            ax.annotate('', xy=(x + box_w + gap - 0.08, 2.05),
                        xytext=(x + box_w + 0.05, 2.05),
                        arrowprops=dict(arrowstyle='->', color='#7F8C8D', lw=1.5))

    ax.text(7, 3.9, '单张图像推理处理流程', ha='center', fontsize=14, fontweight='bold', color='#2C3E50')

    plt.tight_layout()
    path = os.path.join(OUT_DIR, "fig2_inference_pipeline.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def draw_model_comparison():
    """绘制YOLOv8模型规格对比图"""
    fig, axes = plt.subplots(1, 3, figsize=(14, 5))
    fig.patch.set_facecolor('white')

    models = ['n-seg', 's-seg', 'm-seg', 'l-seg', 'x-seg']
    map50 = [0.85, 0.88, 0.91, 0.92, 0.93]
    map50_95 = [0.65, 0.70, 0.74, 0.76, 0.77]
    speed = [10, 18, 35, 50, 60]
    params = [3.4, 11.8, 27.3, 46.0, 71.8]
    sizes = [6.7, 22.5, 52.0, 87.7, 136.8]
    colors_bar = ['#2ECC71', '#3498DB', '#E8913A', '#9B59B6', '#E74C3C']

    # 图1: mAP对比
    ax = axes[0]
    x = np.arange(len(models))
    w = 0.35
    ax.bar(x - w/2, map50, w, label='mAP50', color='#3498DB', edgecolor='white')
    ax.bar(x + w/2, map50_95, w, label='mAP50-95', color='#2ECC71', edgecolor='white')
    ax.set_xticks(x)
    ax.set_xticklabels(models, fontsize=10)
    ax.set_ylabel('mAP', fontsize=11)
    ax.set_title('精度对比 (Accuracy)', fontsize=13, fontweight='bold')
    ax.legend(fontsize=9)
    ax.set_ylim(0.5, 1.0)
    for bar in ax.patches:
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01,
                f'{bar.get_height():.2f}', ha='center', fontsize=8)

    # 图2: 推理速度与参数
    ax = axes[1]
    ax2 = ax.twinx()
    bars = ax.bar(x, speed, 0.5, color=colors_bar, edgecolor='white')
    ax2.plot(x, params, 'D-', color='#E74C3C', linewidth=2, markersize=8, label='参数量(M)')
    ax.set_xticks(x)
    ax.set_xticklabels(models, fontsize=10)
    ax.set_ylabel('推理速度 (ms)', fontsize=11, color='#2C3E50')
    ax2.set_ylabel('参数量 (M)', fontsize=11, color='#E74C3C')
    ax.set_title('速度与规模对比', fontsize=13, fontweight='bold')
    for bar, v in zip(bars, speed):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                f'{v}ms', ha='center', fontsize=8)
    lines1, labels1 = ax.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax.legend(lines1 + lines2, labels1 + labels2, loc='upper left', fontsize=9)

    # 图3: 模型大小
    ax = axes[2]
    bars = ax.barh(models, sizes, color=colors_bar, edgecolor='white')
    ax.set_xlabel('模型大小 (MB)', fontsize=11)
    ax.set_title('存储占用对比', fontsize=13, fontweight='bold')
    for bar, s in zip(bars, sizes):
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2,
                f'{s}MB', va='center', fontsize=9)

    fig.suptitle('YOLOv8-seg 不同规格模型性能对比', fontsize=16, fontweight='bold', y=1.02)
    plt.tight_layout()
    path = os.path.join(OUT_DIR, "fig3_model_comparison.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def draw_scheme_comparison():
    """绘制三方案雷达对比图"""
    fig, ax = plt.subplots(1, 1, figsize=(9, 8), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor('white')

    categories = ['检测精度', '面积计算', '抗干扰能力', '实时性', '部署难度', '标注成本', '可扩展性']
    N = len(categories)

    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]

    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=11)

    values_cv = [2, 3, 1, 5, 5, 5, 1]
    values_detect = [3.5, 1, 4, 4, 3, 4, 3]
    values_seg = [5, 5, 5, 3, 3, 1, 4]
    values_cv += values_cv[:1]
    values_detect += values_detect[:1]
    values_seg += values_seg[:1]

    ax.fill(angles, values_cv, alpha=0.15, color='#E74C3C')
    ax.plot(angles, values_cv, 'o-', linewidth=2, color='#E74C3C', label='传统CV方法')
    ax.fill(angles, values_detect, alpha=0.15, color='#F39C12')
    ax.plot(angles, values_detect, 's-', linewidth=2, color='#F39C12', label='YOLOv8-Detect')
    ax.fill(angles, values_seg, alpha=0.2, color='#2ECC71')
    ax.plot(angles, values_seg, 'D-', linewidth=2.5, color='#2ECC71', label='YOLOv8-Segment ★')

    ax.set_ylim(0, 5.5)
    ax.set_yticks([1, 2, 3, 4, 5])
    ax.set_yticklabels(['差', '较差', '一般', '好', '很好'], fontsize=9)
    ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=11)
    ax.set_title('三种裂缝检测方案综合对比', fontsize=15, fontweight='bold', pad=25)

    plt.tight_layout()
    path = os.path.join(OUT_DIR, "fig4_scheme_radar.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def draw_milestone():
    """绘制技术路线时间线"""
    fig, ax = plt.subplots(1, 1, figsize=(14, 2.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 2.5)
    ax.axis('off')

    milestones = [
        ('数据准备', 1, '#3498DB'),
        ('数据分析', 3, '#3498DB'),
        ('模型训练', 5, '#2ECC71'),
        ('模型评估', 7, '#2ECC71'),
        ('模型优化', 9, '#E67E22'),
        ('模型导出', 11, '#9B59B6'),
        ('部署应用', 13, '#E74C3C'),
    ]

    # 横线
    ax.plot([0.5, 13.8], [1.5, 1.5], '-', color='#BDC3C7', linewidth=3, zorder=1)

    for label, x, c in milestones:
        ax.plot(x, 1.5, 'o', markersize=14, color=c, zorder=2, markeredgecolor='white', markeredgewidth=2)
        ax.text(x, 1.9, label, ha='center', fontsize=10, fontweight='bold', color=c)
        ax.text(x, 0.95, f'阶段{int((x+1)/2)}', ha='center', fontsize=8, color='#7F8C8D')

    # 大标题
    ax.text(7, 0.2, '全流程技术路线', ha='center', fontsize=14, fontweight='bold', color='#2C3E50')

    # 箭头 (用简单方式)
    ax.annotate('→', xy=(0.7, 1.5), fontsize=20, color='#95A5A6')
    for i in range(1, 7):
        ax.annotate('→', xy=(milestones[i][1] - 0.55, 1.5), fontsize=20, color='#95A5A6')

    plt.tight_layout()
    path = os.path.join(OUT_DIR, "fig5_milestone.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


# ═══════════════════════════════════════════════════════════
# 生成 Word 文档
# ═══════════════════════════════════════════════════════════

def create_docx():
    """创建完整方案报告 docx"""
    print("Drawing diagrams...")

    fig1 = draw_system_architecture()
    fig2 = draw_inference_pipeline()
    fig3 = draw_model_comparison()
    fig4 = draw_scheme_comparison()
    fig5 = draw_milestone()

    print("Diagrams done, generating Word doc...")

    doc = Document()

    # ── 页面设置 ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # ── 标题 ──
    title = doc.add_heading('墙面裂缝智能巡检系统 — 实施方案', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    # ═══════════════════════════════════════════════════
    # 第一章：基本技术路线
    # ═══════════════════════════════════════════════════
    doc.add_heading('一、基本技术路线', level=1)

    doc.add_paragraph(
        '本课题采用"数据准备 → 模型训练 → 模型评估 → 模型优化 → 推理部署"的全流程技术路线，'
        '基于YOLOv8实例分割模型实现墙面裂缝的像素级检测与分割。整体系统分为五个层次：数据层、'
        '训练层、推理层、应用层和部署层，各层协同工作，形成完整的裂缝智能巡检解决方案。'
    )

    doc.add_paragraph('系统总体架构如下图所示：')
    doc.add_picture(fig1, width=Inches(5.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('图1  系统总体架构框图').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        '技术路线分为以下七个阶段：\n'
        '阶段一：数据准备 — 整理crack-seg数据集，包含3717张训练图像、112张验证图像和200张测试图像，'
        '统一格式为YOLO多边形分割标注。\n'
        '阶段二：数据分析 — 统计图像尺寸分布、裂缝实例数量、标注多边形点数等，为训练参数选择提供依据。\n'
        '阶段三：模型训练 — 加载YOLOv8n-seg预训练权重，在crack-seg数据集上训练100轮，采用混合精度训练(AMP)'
        '加速并降低显存占用，同时启用早停机制防止过拟合。\n'
        '阶段四：模型评估 — 在验证集和测试集上计算mAP50、mAP50-95、Precision、Recall等指标，绘制训练损失'
        '曲线和mAP曲线，为模型选型提供数据支持。\n'
        '阶段五：模型优化 — 对比不同规格模型(n/s/m/l/x)的性能，选择最佳精度-速度平衡点。\n'
        '阶段六：模型导出 — 将训练好的PyTorch模型导出为ONNX、TensorRT、OpenVINO等工业部署格式。\n'
        '阶段七：部署应用 — 通过CLI命令行、Gradio Web界面、Docker容器等方式部署系统。'
    )

    doc.add_picture(fig5, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('图2  全流程技术路线').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 核心推理流程 ──
    doc.add_heading('1.1 核心推理流程', level=2)
    doc.add_paragraph(
        '单张墙面图像经过预处理、特征提取、多尺度融合、分割头预测和后处理，最终输出裂缝的边界框、'
        '像素级分割掩码和置信度。基于掩码面积占图像总像素的比例，系统自动将裂缝严重程度分为三级：'
        '轻微(≤2%)、中等(2%~5%)和严重(>5%)。'
    )

    doc.add_picture(fig2, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('图3  单张图像推理处理流程').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ═══════════════════════════════════════════════════
    # 第二章：核心设计方案
    # ═══════════════════════════════════════════════════
    doc.add_heading('二、核心设计方案及对比', level=1)

    doc.add_paragraph(
        '裂缝检测是本课题的核心任务，选择合适的检测算法直接决定系统的精度、速度和适用性。'
        '我们调研了三类主流方法：传统计算机视觉方法、YOLOv8目标检测方法和YOLOv8实例分割方法。'
        '以下从多个维度进行详细对比，并结合课题需求给出选型结论。'
    )

    # 2.1 传统CV
    doc.add_heading('2.1 方案一：传统计算机视觉方法', level=2)
    doc.add_paragraph(
        '基于OpenCV实现，核心流程为：灰度化 → 高斯滤波去噪 → 自适应阈值二值化(解决墙面亮度不均) '
        '→ Canny边缘检测 → 位与融合(结合阈值与边缘结果) → 形态学闭运算(连接裂缝断点) → '
        '轮廓查找 → 面积过滤(剔除面积<50像素的噪声)。'
    )
    doc.add_paragraph(
        '优点：无需训练数据和GPU，CPU即可实时运行，每一步可解释，适合快速验证。\n'
        '缺点：对光照和纹理变化敏感，阈值需手动调参，无法区分裂缝与墙面伪影（阴影、污渍等），'
        '难以精确计算裂缝面积占比和严重程度。'
    )

    # 2.2 YOLOv8 Detect
    doc.add_heading('2.2 方案二：YOLOv8目标检测方法', level=2)
    doc.add_paragraph(
        '使用YOLOv8的检测模式(Detect)，仅输出裂缝的矩形边界框(Bounding Box)，不做像素级分割。'
        '标注格式为矩形框，标注成本相对较低。'
    )
    doc.add_paragraph(
        '优点：能定位裂缝位置，推理速度快，标注成本低。\n'
        '缺点：无法精确计算裂缝面积(矩形框包围区域包含大量非裂缝背景)，'
        '无法评估裂缝严重程度，对细长弯曲裂缝的边界框拟合不精确。'
    )

    # 2.3 YOLOv8 Segment
    doc.add_heading('2.3 方案三：YOLOv8实例分割方法（选定方案）', level=2)
    doc.add_paragraph(
        '使用YOLOv8的分割模式(Segment)，同时输出裂缝的边界框和像素级分割掩码(Mask)。'
        '基于掩码面积可精确计算裂缝面积占比，并自动分级严重程度：\n'
        '  · 🟢 轻微：裂缝面积占比 ≤ 2%\n'
        '  · 🟡 中等：2% < 裂缝面积占比 ≤ 5%\n'
        '  · 🔴 严重：裂缝面积占比 > 5%'
    )
    doc.add_paragraph(
        '优点：像素级精确提取裂缝区域，可计算裂缝面积与严重程度，深度学习特征具有强抗干扰能力，'
        '支持n/s/m/l/x五种规模按需选择。\n'
        '缺点：需多边形标注(标注成本高)，推理速度略慢于目标检测模型。'
    )

    # 雷达图
    doc.add_picture(fig4, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('图4  三种方案综合对比雷达图').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2.4 模型规格对比
    doc.add_heading('2.4 YOLOv8-seg 模型规格对比', level=2)
    doc.add_paragraph(
        'YOLOv8-seg提供nano/small/medium/large/xlarge五种规格，下表给出在同一crack-seg数据集上的'
        '预期性能指标：'
    )

    # 表格
    table = doc.add_table(rows=6, cols=7, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['模型', '参数量', 'mAP50', 'mAP50-95', '推理速度(ms)', '模型大小(MB)', '适用场景']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    data = [
        ['YOLOv8n-seg', '3.4M', '~0.85', '~0.65', '~10', '6.7', '边缘设备、实时检测'],
        ['YOLOv8s-seg', '11.8M', '~0.88', '~0.70', '~18', '22.5', '平衡精度与速度'],
        ['YOLOv8m-seg', '27.3M', '~0.91', '~0.74', '~35', '52.0', '常规部署'],
        ['YOLOv8l-seg', '46.0M', '~0.92', '~0.76', '~50', '87.7', '高精度场景'],
        ['YOLOv8x-seg', '71.8M', '~0.93', '~0.77', '~60', '136.8', '离线分析、研究'],
    ]
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph('表1  YOLOv8-seg 各规格模型性能对比').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_picture(fig3, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('图5  模型规格性能对比图').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2.5 方案总结对比
    doc.add_heading('2.5 方案综合对比与选型结论', level=2)

    table2 = doc.add_table(rows=8, cols=5, style='Light Grid Accent 1')
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers2 = ['对比维度', '传统CV方法', 'YOLOv8-Detect', 'YOLOv8-Segment ★', '说明']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    comparison_data = [
        ['检测精度', '★★☆', '★★★☆', '★★★★★', '分割 > 检测 > 传统'],
        ['面积计算', '近似(像素计数)', '不支持', '精确(掩码级)', '关键需求'],
        ['抗干扰能力', '弱', '强', '强', '深度学习优势'],
        ['标注成本', '无', '低(矩形框)', '高(多边形)', '需考虑数据集'],
        ['实时性', '★★★★★', '★★★★☆', '★★★☆☆', '可硬件加速'],
        ['部署难度', '简单', '中', '中', '多格式导出'],
        ['严重程度分级', '不支持', '不支持', '支持', '核心需求'],
    ]
    for r, row_data in enumerate(comparison_data):
        for c, val in enumerate(row_data):
            cell = table2.rows[r+1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph('表2  三种方案综合对比').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    doc.add_paragraph(
        '★ 最终选型结论：采用 YOLOv8-Segment（实例分割）作为主方案。理由如下：\n\n'
        '1. 需求匹配：课题要求裂缝面积计算与严重程度分级，只有实例分割能提供像素级掩码，'
        '这是目标检测方法无法实现的。\n\n'
        '2. 精度优势：深度特征学习对墙面纹理、光照变化具有天然鲁棒性，'
        'mAP50可达0.85以上，远优于传统CV方法的启发式阈值分割。\n\n'
        '3. 工程灵活性：支持n/s/m/l/x五种规模按需切换，训练部署一体化，'
        '通过ONNX/TensorRT/OpenVINO等导出格式可适配GPU服务器、边缘设备和移动端等多种场景。\n\n'
        '4. 可扩展性：在YOLOv8框架下可轻松扩展为多类别检测（如增加脱落、渗水等缺陷类别），'
        '也可接入目标跟踪实现裂缝变化监测。\n\n'
        '同时保留传统CV方法(cvmain.py)作为辅助方案，用于无需GPU场景下的基础检测和快速对比验证。'
    )

    # 保存
    docx_path = os.path.join(OUT_DIR, "实施方案_墙面裂缝智能巡检系统.docx")
    doc.save(docx_path)
    print(f"\nReport saved: {docx_path}")
    return docx_path


if __name__ == '__main__':
    path = create_docx()
    print(f"文件路径: {os.path.abspath(path)}")
